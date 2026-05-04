"""Generate a Word document explaining the Sapling engine's macadamia logic.

Run from repo root:
    sapling-api/venv/bin/python scripts/generate_macadamia_engine_doc.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


SAPLING_GREEN = RGBColor(0x2F, 0x6B, 0x3A)
SLATE = RGBColor(0x33, 0x3F, 0x48)
MUTED = RGBColor(0x6A, 0x73, 0x80)


def _style_normal(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = SLATE


def _h1(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = SAPLING_GREEN
    p.paragraph_format.space_after = Pt(6)


def _h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = SAPLING_GREEN
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)


def _h3(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = SLATE
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)


def _para(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)


def _meta(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def _bullet(doc: Document, text: str, level: int = 0) -> None:
    style = "List Bullet" if level == 0 else "List Bullet 2"
    doc.add_paragraph(text, style=style)


def _table(doc: Document, header: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = ""
        para = hdr[i].paragraphs[0]
        run = para.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            para = cells[i].paragraphs[0]
            run = para.add_run(val)
            run.font.size = Pt(10)


def build(output_path: Path) -> None:
    doc = Document()
    _style_normal(doc)

    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # ── Cover ──────────────────────────────────────────────────────────
    _h1(doc, "Sapling Programme Engine — Macadamia")
    _meta(doc, "How the engine builds a macadamia programme: factors, sources, drivers")
    _meta(doc, "Production engine on commit 2f3d6f8 · app.saplingfertilizer.co.za")

    _para(
        doc,
        "This document describes what the Sapling programme engine takes into account when "
        "building a macadamia programme — every input it consumes, every reasoning module it "
        "runs, every cited source it pulls from, and which of those factors do the heavy "
        "lifting in the final recommendation. It reflects the engine actually deployed in "
        "production, not a roadmap."
    )

    # ── 1. Inputs ──────────────────────────────────────────────────────
    _h2(doc, "1. Inputs the engine consumes")

    _h3(doc, "Per-block inputs (from the wizard)")
    _bullet(doc, "Block area (ha) — drives kg/ha → total kg conversion and area-weighted aggregation when blocks are clustered.")
    _bullet(doc, "Soil analysis: 30 canonical parameters normalised by the soil_canonicaliser (pH, P-Bray/Olsen/Citric, exchangeable K/Ca/Mg/Na, CEC, exchangeable Al, Organic C, Total N, EC, S, Zn, B, Cu, Mn, Fe, Mo).")
    _bullet(doc, "Lab name + lab method (Mehlich-3, Bray-1, Ambic-2, Citric, Olsen) — drives sufficiency-threshold selection and any unit-suffix corrections.")
    _bullet(doc, "Sample date and sample ID — used for trend analysis when more than one analysis exists per block.")
    _bullet(doc, "Tree population (pop_per_ha) — e.g. Levubu 312/ha, White River 555/ha, Barberton 6×3 m. The engine scales the per-ha kg/ha target by block_pop_per_ha ÷ crop_requirements.pop_per_ha (SAMAC Schoeman 2021 density convention).")
    _bullet(doc, "Tree age — drives the perennial_age_factors curve (young/non-bearing trees scaled down vs full-bearing reference).")
    _bullet(doc, "Pre-season inputs already applied (kraal manure, dolomitic lime, gypsum, RP) with date + rate — used by the pre-season module to compute residuals to subtract from in-season targets.")
    _bullet(doc, "Mid-season leaf analysis (Season Tracker re-entry only) — triggers leaf-correction foliar events.")

    _h3(doc, "Programme-level inputs")
    _bullet(doc, "Crop = Macadamia (canonicalised — accepts \"mac\", \"macs\", \"makadamia\", \"macadamia nuts\", \"noote\").")
    _bullet(doc, "Yield target (t/ha DIS — Dry-In-Shell). Defaults to crop_requirements.default_yield (full-bearing potential) if not provided, with an Assumption row surfaced.")
    _bullet(doc, "Planting date (orchard establishment) — anchors the growth-stage calendar.")
    _bullet(doc, "Build date — drives pre-season Mode A/B/C decision (lead time available, residual subtraction, or lost opportunity).")
    _bullet(doc, "Method availability — five flags: dry_broadcast, dry_band, dry_side_dress, fertigation, foliar_sprayer. Without a foliar sprayer, foliar B/Zn cannot fire and are downgraded to a warning rather than a recommendation.")
    _bullet(doc, "Application months chosen by the agronomist in the wizard — owned by the agronomist (house rule: engine never overrides a chosen month for nutrient correctness).")
    _bullet(doc, "Cluster method overrides per zone — single-block programmes resolve through a block_to_cluster map so foliar-excluded blocks do not silently get foliar events.")

    # ── 2. Engine pipeline ─────────────────────────────────────────────
    _h2(doc, "2. Engine pipeline — modules in order")

    _para(
        doc,
        "The orchestrator (app/services/programme_builder_orchestrator.py) chains seven "
        "reasoning modules, each carrying its own provenance. Every number that lands in the "
        "final artifact carries a SourceCitation (source_id + section + tier 1–6)."
    )

    _h3(doc, "2.1 Soil canonicaliser")
    _para(
        doc,
        "Single source of truth across 30 canonical soil parameters. Resolves unit-suffix "
        "ambiguity and parameter-name mismatches between labs. Fixes the 374% ESP class of "
        "bug at the front door."
    )

    _h3(doc, "2.2 Soil factor reasoner")
    _para(
        doc,
        "Reads the canonical soil values and produces structured findings. Inputs the engine "
        "evaluates for macadamia:"
    )
    _bullet(doc, "pH (water + KCl) — low-pH → lime recommendation; high-pH → elemental sulphur. Crop-specific thresholds for macadamia via crop_sufficiency_overrides.")
    _bullet(doc, "Ca:Mg ratio — low Ca:Mg at acceptable pH → gypsum (no Al/SAR required).")
    _bullet(doc, "Al saturation (exchangeable Al ÷ ECEC × 100) — toxic-Al trigger pushes N-source to Ca-Nitrate and lime priority up.")
    _bullet(doc, "ESP (exchangeable sodium percentage) — ≥ 15% sodic, 10–15% trending.")
    _bullet(doc, "SAR — non-sodic < 13, moderate 13–20, sodic > 20 (USSL Handbook 60).")
    _bullet(doc, "EC — > 2 mS/cm topsoil flags salt-sensitive crop risk.")
    _bullet(doc, "P:Zn ratio — > 150:1 triggers foliar Zn (high P locks up soil Zn). Foth & Ellis 1997.")
    _bullet(doc, "Ca:B ratio — > 1000:1 triggers foliar B on high-Ca soils. Shorrocks 1997.")
    _bullet(doc, "Fe:Mn and Cu:Mo antagonisms — flagged for narrative.")
    _bullet(doc, "Irrigation water quality (when supplied): EC, SAR, RSC, HCO3 — FAO Ayers & Westcot 1985.")

    _h3(doc, "2.3 Pre-season module")
    _para(doc, "Three modes — not mutually exclusive:")
    _bullet(doc, "Mode A — Recommend: lime / gypsum / RP / S° with apply-by date computed from build date and reaction-time profile (lime 3–6 mo, dolomitic 4–8 mo, gypsum 0.5–2 mo, RP 12–36 mo).")
    _bullet(doc, "Mode B — Subtract residual: applied_amount × reaction_pct × soil_fixation_factor → kg/ha already in the soil at planting. P fixation 50% on high-Al, 85% otherwise. N leaching 10% wet / 50% dry season.")
    _bullet(doc, "Mode C — Lost opportunity: insufficient lead time → OutstandingItems documenting what could have been done with more notice + RiskFlags about consequences of going ahead.")

    _h3(doc, "2.4 Target computation (per nutrient)")
    _para(doc, "Strict precedence order:")
    _bullet(doc, "Tier 1–2 rate-table hit (FERTASA fertilizer_rate_tables) — 908 rows total, sparse for macadamia. When present, returns Rate_Mid = (rate_min + rate_max)/2 with full citation.")
    _bullet(doc, "Tier 1 cation-ratio path for Ca/Mg only — FERTASA §5.2.2 base-saturation when CEC ≥ 3 cmol/kg. Skipped on acid-loving crops via crop_calc_flags (not flagged for macadamia).")
    _bullet(doc, "Tier 6 heuristic = per_unit × yield_target × adjustment_factor. per_unit from crop_requirements (per-ton removal coefficients). adjustment_factor from the per-nutrient-group curves (P / K / ca_mg / N / S / micros — see §3 below).")
    _bullet(doc, "Tier 6 unadjusted fallback when soil_values is missing the parameter — surfaced as calc_path=\"unadjusted\" + an OutstandingItem so the agronomist sees the recommendation is uninformed by soil state.")
    _para(doc, "After per-nutrient compute, two scaling steps fire for macadamia (perennial):")
    _bullet(doc, "Density scaling: target_kg_per_ha × (block_pop_per_ha ÷ reference_pop_per_ha) — SAMAC convention.")
    _bullet(doc, "Age scaling: N/P/K × perennial_age_factors[bracket] for tree_age — young trees get fractional rates.")
    _para(doc, "P → P2O5 (×2.291) and K → K2O (×1.205) before downstream modules consume it.")

    _h3(doc, "2.5 Stage splitter")
    _para(
        doc,
        "Apportions the residual-adjusted season targets across crop stages using SAMAC "
        "Schoeman 2021 weights for macadamia (Tier 1, all 4 stages confirmed). The shape "
        "biases the K peak into Oct–Dec (nut development + oil accumulation):"
    )
    _table(
        doc,
        ["Stage", "SA Months", "N %", "P %", "K %", "Ca %", "Mg %"],
        [
            ["1 — Post-harvest / dormant",      "Apr–Jun", "10",  "15", "10", "15", "15"],
            ["2 — Flower init / pre-flower",    "Jul",     "20",  "25", "15", "30", "25"],
            ["3 — Flowering + nut set",         "Aug–Sep", "30",  "30", "40", "30", "30"],
            ["4 — Nut development / oil fill",  "Oct–Dec", "25",  "20", "30", "20", "20"],
            ["5 — Pre-harvest taper",           "Jan–Mar", "15",  "10",  "5",  "5", "10"],
        ],
    )
    _bullet(doc, "K:N ratio rises from 1:1 in young trees to 1.25–1.5:1 on bearing trees.")
    _bullet(doc, "P total ≈ N/5 (FERTASA 5.8.1 explicit).")
    _bullet(doc, "FERTASA constraint: N supplemented Apr–Dec only — vegetative growth Jan–Mar hampers nut growth and oil accumulation.")

    _h3(doc, "2.6 Method selector")
    _para(
        doc,
        "Per (stage × nutrient) decides one primary method given availability + nutrient "
        "behaviour. Macadamia routing tendencies:"
    )
    _bullet(doc, "N — fertigation when drip available, otherwise dry-broadcast or banded LAN/Urea split across the SAMAC calendar.")
    _bullet(doc, "P — dry-broadcast/band at post-harvest or via fertigation if soluble P (MAP/MKP) suits the stream.")
    _bullet(doc, "K — fertigation preferred during the Oct–Dec peak; dry-broadcast KCl/SOP otherwise. Selection respects FERTASA Part-A/Part-B stream rules in the consolidator.")
    _bullet(doc, "Ca — dry lime/gypsum at pre-season; Ca-Nitrate fertigation when soil Al is flagged.")
    _bullet(doc, "Zn / B — foliar by default for macadamia regardless of soil status (stage_peak rules — see §2.7).")
    _bullet(doc, "Foliar single-event caps respected: N ≤ 5, K ≤ 2.5, B ≤ 0.4, Zn ≤ 1.0 kg/ha per pass. Excess routes to root.")

    _h3(doc, "2.7 Foliar trigger engine")
    _para(doc, "Foliar is contingent — at least one of four triggers must fire:")
    _bullet(doc, "soil_availability_gap — e.g. P:Zn > 150:1 or Ca:B > 1000:1 from the soil reasoner.")
    _bullet(doc, "stage_peak_demand — pre-flowering B and vegetative-flush Zn for macadamia (SAMAC Schoeman 2021, Tier 1).")
    _bullet(doc, "quality_window — bitter-pit-style Ca windows (not standard for mac; cited for completeness).")
    _bullet(doc, "leaf_correction — mid-season leaf deficiency from Season Tracker re-entry.")
    _para(doc, "Macadamia stage_peak rules baked in:")
    _table(
        doc,
        ["Trigger", "Stage", "Product", "Rate", "Source", "Tier"],
        [
            ["Pre-flowering B",     "Week 28 — pre-bloom",       "Solubor (20.5% B)",       "0.1% spray",   "SAMAC Schoeman 2021 §3.2", "1"],
            ["Vegetative-flush Zn", "Week 32 — leaf expansion",  "Zn Sulphate (36% Zn)",    "0.5% spray",   "SAMAC Schoeman 2021 §3.2", "1"],
        ],
    )
    _para(doc, "If has_foliar_sprayer = false, these become RiskFlags (\"foliar trigger fired but no sprayer available\"), not events — agronomist decides.")

    _h3(doc, "2.8 Consolidator + similarity merger")
    _para(
        doc,
        "Groups per-stage method assignments into the smallest set of distinct blends without "
        "mixing incompatible streams. Macadamia constraints honoured:"
    )
    _bullet(doc, "Every dry blend ≥ 50% organic carrier (Sapling house rule — hard constraint, narrated in the artifact).")
    _bullet(doc, "FERTASA Part-A / Part-B fertigation streams kept separate — Ca-Nitrate (Part B) never co-applied with sulphate or phosphate sources.")
    _bullet(doc, "Nutrient-antagonism timing walls — citrus-style FERTASA 5.7.3 N+K split logic where applicable.")
    _bullet(doc, "Compound-product preference where one product covers ≥ 60% of the stage demand.")
    _bullet(doc, "Similarity merger collapses near-identical blends across blocks within a cluster (NPK ratio L1 < 0.15 margin).")

    _h3(doc, "2.9 Liquid optimizer (MILP)")
    _para(
        doc,
        "When the consolidator hands off a fertigation event, the liquid_optimizer (MILP "
        "with pair-incompat constraints) sizes the Part-A / Part-B tank mixes in m/m or g/L "
        "notation. Already mature; powers Quick Blend and the programme builder's "
        "fertigation step."
    )

    _h3(doc, "2.10 Blend validator")
    _para(doc, "Per-event ceilings warn (not block) when a blend overshoots realistic application rates:")
    _bullet(doc, "Dry-broadcast 1500 kg/ha")
    _bullet(doc, "Side-dress 600 kg/ha")
    _bullet(doc, "Fertigation 500 L/ha")
    _bullet(doc, "Foliar 800 L/ha")
    _para(doc, "Tunable via the application_rate_limits DB table.")

    _h3(doc, "2.11 Risk flags + outstanding items")
    _para(
        doc,
        "Cross-cutting consequences surface here: missing soil parameters, low data confidence, "
        "agronomist-action items (\"order Solubor by week 26\"), pre-season lost opportunities, "
        "and tier-rollup data confidence (worst tier across all nutrients drives the artifact's "
        "data-confidence band)."
    )

    # ── 3. Adjustment factor curves ───────────────────────────────────
    _h2(doc, "3. Adjustment factor curves (the per-nutrient drawdown logic)")
    _para(
        doc,
        "Migration 077 split a single \"cations\" curve into six per-nutrient-group curves. "
        "These are the multipliers applied to per_unit × yield in the heuristic path:"
    )
    _table(
        doc,
        ["Soil class", "P", "K", "Ca/Mg", "N", "S", "Micros"],
        [
            ["Very Low",  "1.50", "1.50", "1.25", "1.00", "1.00", "1.00"],
            ["Low",       "1.25", "1.25", "1.10", "1.00", "1.00", "1.00"],
            ["Optimal",   "1.00", "1.00", "1.00", "1.00", "1.00", "1.00"],
            ["High",      "0.75", "0.50", "0.75", "1.00", "1.00", "1.00"],
            ["Very High", "0.50", "0.00", "0.50", "1.00", "1.00", "1.00"],
        ],
    )
    _bullet(doc, "P drawdown moderate (sticky soil P = real reservoir).")
    _bullet(doc, "K drawdown aggressive — luxury K consumption is real, K is mobile, soil-K = 0 at Very High is correct.")
    _bullet(doc, "Ca/Mg drawdown gentle to preserve base saturation.")
    _bullet(doc, "N flat — soil mineral-N test is not a reliable scaling input.")
    _bullet(doc, "S flat — sulphate leaches seasonally; classification drawdown does not apply (S_pct minima in blends handle replenishment).")
    _bullet(doc, "Micros flat — correction is foliar, not base-rate scaling.")

    # ── 4. Sources ────────────────────────────────────────────────────
    _h2(doc, "4. Sources the engine pulls from for macadamia")

    _h3(doc, "Tier 1 — SA authoritative (preferred)")
    _bullet(doc, "FERTASA Handbook ch. 5.8.1 Macadamia (Hawksworth & Sheard 2022) — N calendar, K peak Oct–Dec, P=N/5 ratio, K:N progression by tree age. The single most-cited source in the macadamia path.")
    _bullet(doc, "SAMAC Schoeman 2021 — regional K rates (North Coast 100–150, White River 130–150, Nelspruit/Levubu 200–250, Barberton 250–300 kg K/ha for cv 695 micro-irr); cultivar leaf-norm offsets (A4 / N2 / 695 / 814 / 816); per-ton DIS removal (kg/4t DIS: N 80, P 8, K 80, Ca 70, Mg 12, Zn 2, Cu 0.3, B 2, Mn 4, Fe 4); pre-flower B and vegetative Zn foliar calendar.")
    _bullet(doc, "SAMAC Information Hub — supplementary cultivar / agronomy guidance.")

    _h3(doc, "Tier 2 — international parallels (where SA is silent)")
    _bullet(doc, "Foth & Ellis Soil Fertility 2nd ed. (1997) p.224 — P:Zn 150:1 antagonism threshold.")
    _bullet(doc, "Shorrocks (1997) Plant & Soil 193:121–148 — Ca:B 1000:1 induced-deficiency threshold.")
    _bullet(doc, "USSL Agriculture Handbook 60 — SAR + ESP classification.")
    _bullet(doc, "FAO Irrigation & Drainage Paper 29 (Ayers & Westcot 1985) — water-quality thresholds.")
    _bullet(doc, "IFA World Fertilizer Use Manual 1992 — N-acidification CaO equivalents (lime-from-N reaction-time profiles).")

    _h3(doc, "Tier 3 — peer-reviewed papers")
    _bullet(doc, "Manson & Sheard (2007) — KZN/Cedara macadamia formulas: N g/tree/year by age (Y1=50 → mature=360 g/tree), K g/tree/year, P = (target − soil P) × PRF capped 50 kg/ha, lime formula, Mg formula for soil Mg < 100 mg/L. Cedara Ambic-2 extraction. Plug-through validation reference.")
    _bullet(doc, "Du Preez (2014) Stellenbosch MSc Beaumont multi-site trial — White River 555/ha 6×3 m Hutton ~35% clay 4.4 t/ha biennial DIS; Barberton 6-y.o. Beaumont; 36-orchard historical 2010–11; crop-load classes low ≤3500 / med 3500–4900 / high ≥4900 kg/ha; leaf-element/yield index regression coefficients.")
    _bullet(doc, "Wolstenholme — Macadamia Fertilization (FERTASA Symposium 2017) — N split 25% Jul / 40% Sep / 25% Nov / 10% Jan supplementary.")

    _h3(doc, "Tier 6 — implementer convention (visibly de-rated in the artifact)")
    _bullet(doc, "Standard SA agronomy + Foth-Ellis + IFA where no specific handbook citation exists. The renderer surfaces the worst tier across all rows as the artifact's data-confidence band.")

    # ── 5. Biggest drivers ─────────────────────────────────────────────
    _h2(doc, "5. The biggest driving factors — what actually moves the recommendation")

    _para(
        doc,
        "Across a representative macadamia programme build, these are the inputs whose values "
        "swing the final blend the most. Listed roughly in order of magnitude."
    )

    _h3(doc, "1. Soil-test classification per nutrient")
    _para(
        doc,
        "The single largest multiplier. K classified Very High zeroes the K target outright "
        "(adjustment factor 0.0); Very Low pushes it to 1.5× the heuristic. The same curve "
        "applied to P or Ca/Mg can swing the recommendation by ±50%. This is why a current, "
        "lab-method-tagged soil analysis is the highest-leverage data point in the build."
    )

    _h3(doc, "2. Tree population (pop_per_ha) — density scaling")
    _para(
        doc,
        "Levubu 312/ha vs White River 555/ha is a ~78% per-ha target difference for the same "
        "kg-per-tree intent. The engine scales by block_pop_per_ha ÷ reference_pop_per_ha "
        "(SAMAC Schoeman 2021 convention) — an entered density is one of the two largest "
        "swings in the per-ha rate."
    )

    _h3(doc, "3. Tree age via perennial_age_factors")
    _para(
        doc,
        "Young (Y1–Y4) blocks scale N/P/K to a fraction of full-bearing rates. A misentered "
        "tree age on a 6-year-old block can over-fertilize by 60%+ or starve a mature block "
        "by the same. The engine never assumes — it uses whatever tree_age the agronomist "
        "supplied."
    )

    _h3(doc, "4. Yield target × per-ton DIS removal")
    _para(
        doc,
        "Multiplicative across every nutrient. SAMAC Schoeman per-4-t-DIS removal: N 80, P 8, "
        "K 80, Ca 70, Mg 12, Zn 2, B 2 kg. A 3.5 t/ha vs 5.0 t/ha expectation moves the season "
        "target proportionally."
    )

    _h3(doc, "5. Region / cultivar (via yield + K-rate band)")
    _para(
        doc,
        "Schoeman's regional K bands (North Coast 100–150 → Barberton 250–300 kg K/ha) reflect "
        "real climate × soil × cultivar interactions. The agronomist's choice of yield target "
        "implicitly encodes this; A4 / N2 / 695 / 814 / 816 leaf-norm offsets feed the leaf-correction "
        "path when Season Tracker re-enters mid-season."
    )

    _h3(doc, "6. Stage 3 (Aug–Sep flowering + nut set) and Stage 4 (Oct–Dec nut development)")
    _para(
        doc,
        "Together these stages absorb roughly 70% of annual K and 50% of annual N per the "
        "SAMAC Schoeman shape. Any change to these two stage windows propagates almost "
        "directly to the blend recommendation."
    )

    _h3(doc, "7. Soil P:Zn and Ca:B antagonism flags")
    _para(
        doc,
        "On most SA macadamia soils these flags fire, which is why the engine produces a "
        "foliar Zn (vegetative flush) and foliar B (pre-flower) regardless of soil-test "
        "level on those micros. Removing those triggers (e.g. via a cluster method override "
        "that excludes foliar) shifts the whole micro story to soil application + warning flags."
    )

    _h3(doc, "8. Pre-season residuals (manure / lime / RP)")
    _para(
        doc,
        "Mode-B residual subtraction can remove 20–60 kg N/ha and 30–80 kg P2O5/ha from the "
        "in-season target depending on rate, reaction time, and Al-driven fixation factor. "
        "Forgetting to enter a 30 t/ha kraal manure application is a meaningful miss."
    )

    _h3(doc, "9. Method availability (foliar sprayer in particular)")
    _para(
        doc,
        "Without a foliar sprayer the engine cannot fire the SAMAC pre-flower B and "
        "vegetative-flush Zn — these become RiskFlags rather than events. On most SA "
        "soils that means a real micronutrient gap the orchard owner needs to close."
    )

    _h3(doc, "10. Agronomist-chosen application months")
    _para(
        doc,
        "The agronomist owns timing — the engine consolidates events into blends but never "
        "moves a chosen month for nutrient correctness. Bunching all nutrients into one "
        "month (instead of spreading across the SAMAC calendar) collapses the stage shape and "
        "produces a single big blend rather than the 4–5 stage-tuned blends FERTASA implies."
    )

    # ── 6. What the engine does NOT do ─────────────────────────────────
    _h2(doc, "6. What the engine deliberately does not do (today)")
    _bullet(doc, "Does not pick application months. Agronomist owns timing.")
    _bullet(doc, "Does not include cost / cheapest-blend / R-rollups in the programme. Programme builder is agronomy-only; cost lives downstream in /quotes.")
    _bullet(doc, "Does not invent rates. Every number in the artifact has a SourceCitation; missing data surfaces as Tier 6 unadjusted + an OutstandingItem rather than a fabricated number.")
    _bullet(doc, "Does not substitute foliar for unavailable equipment. Triggers become RiskFlags, not silent omissions.")
    _bullet(doc, "Does not yet apply per-cultivar offsets in the heuristic path (A4 vs 695 vs 814) — leaf norms exist, full cultivar-specific adjustment factor work is queued post-Muller-demo.")

    # ── 7. Footer ──────────────────────────────────────────────────────
    _h2(doc, "7. Where to look in the codebase")
    _bullet(doc, "Orchestrator: app/services/programme_builder_orchestrator.py")
    _bullet(doc, "Soil canonicaliser: app/services/soil_canonicaliser.py")
    _bullet(doc, "Soil factor reasoner: app/services/soil_factor_reasoner.py")
    _bullet(doc, "Pre-season module: app/services/pre_season_module.py")
    _bullet(doc, "Target computation: app/services/target_computation.py")
    _bullet(doc, "Stage splitter (macadamia weights): app/services/stage_splitter.py — STAGE_SHAPES[\"Macadamia\"]")
    _bullet(doc, "Method selector: app/services/method_selector.py")
    _bullet(doc, "Foliar trigger engine (macadamia rules): app/services/foliar_trigger_engine.py — STAGE_PEAK_RULES")
    _bullet(doc, "Consolidator (organic-carrier rule, FERTASA streams): app/services/consolidator.py")
    _bullet(doc, "Liquid optimizer: app/services/liquid_optimizer.py")
    _bullet(doc, "Blend validator (rate ceilings): app/services/blend_validator.py")
    _bullet(doc, "Soil engine adjustment factors (migration 077 curves): app/services/soil_engine.py + DB tables adjustment_factors / crop_sufficiency_overrides / crop_calc_flags")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"wrote {output_path}")


if __name__ == "__main__":
    here = Path(__file__).resolve()
    repo_root = here.parent.parent
    out = repo_root / "docs" / "macadamia_engine_overview.docx"
    build(out)
